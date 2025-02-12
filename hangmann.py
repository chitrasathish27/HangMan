from tkinter import *
from tkinter import messagebox
from string import ascii_uppercase
import random
import docx
write=docx.Document()

window = Tk()
window.title('Hangman-Guess Cities or Animal')
word_list= ['MUMBAI','DELHI','BANGLORE','HYDRABAD','AHMEDABAD','CHENNAI','KOLKATA','SURAT','PUNE','JAIPUR','AMRITSAR','ALLAHABAD','RANCHI',
            'LUCKNOW','KANPUR','NAGPUR','INDORE','THANE','BHOPAL','PATNA','GHAZIABAD','AGRA','FARIDABAD','MEERUT','RAJKOT','VARANASI','SRINAGAR',
            'RAIPUR','KOTA','JHANSI','GORILLA','ELEPHANT','TIGER','LION','LEOPARD','DONKEY','JACKAL','GIRAFFE','COW','SHEEP','CHICKEN']
            
photos = [PhotoImage(file="images/hang0.png"), PhotoImage(file="images/hang1.png"), PhotoImage(file="images/hang2.png"),
PhotoImage(file="images/hang3.png"), PhotoImage(file="images/hang4.png"), PhotoImage(file="images/hang5.png"),
PhotoImage(file="images/hang6.png"), PhotoImage(file="images/hang7.png"), PhotoImage(file="images/hang8.png"),
PhotoImage(file="images/hang9.png"), PhotoImage(file="images/hang10.png"), PhotoImage(file="images/hang11.png")]






def newGame():
    global the_word_withSpaces
    global numberOfGuesses
    numberOfGuesses =0
    
    the_word=random.choice(word_list)
    the_word_withSpaces = " ".join(the_word)
    lblWord.set(' '.join("_"*len(the_word)))

def guess(letter):
	global numberOfGuesses
	if numberOfGuesses<11:	
		txt = list(the_word_withSpaces)
		guessed = list(lblWord.get())
		if the_word_withSpaces.count(letter)>0:
			for c in range(len(txt)):
				if txt[c]==letter:
					guessed[c]=letter
				lblWord.set("".join(guessed))
				if lblWord.get()==the_word_withSpaces:
					messagebox.showinfo("Hangman","You guessed it!")
		else:
			numberOfGuesses += 1
			imgLabel.config(image=photos[numberOfGuesses])
			if numberOfGuesses==11:
					messagebox.showwarning("Hangman","Game Over\nThe word is: "+the_word_withSpaces)
def add():
	print(word_list)
	n=input("Enter the new word")
	write.add_paragraph(n)
	write.save("hangmanwords.docx")
	doc = docx.Document("hangmanwords.docx")
	para=doc.paragraphs
	for read in para:
		print(read.text)
		word_list.append(read.text)
	print(word_list)
	exit(0)
def update():
	docc=docx.Document("hangmanwords.docx")
	para=docc.paragraphs
	for read in para:
		print(read.text)
		word_list.append(read.text)
	docc.save("hangmanwords.docx")


imgLabel=Label(window)
imgLabel.grid(row=0, column=0, columnspan=3, padx=10, pady=40)


  
lblWord = StringVar()
Label(window, textvariable  =lblWord,font=('consolas 24 bold')).grid(row=0, column=3 ,columnspan=6,padx=10)

n=0
for c in ascii_uppercase:
    Button(window, text=c, command=lambda c=c: guess(c), font=('Helvetica 18'), width=4).grid(row=1+n//9,column=n%9)
    n+=1

Button(window, text="New\nGame", command=lambda:newGame(), font=("Helvetica 10 bold")).grid(row=3, column=8)
Button(window, text="Add word", command=lambda:add(), font=("Helvetica 10 bold")).grid(row=4, column=4)

update()
newGame()
window.mainloop()
